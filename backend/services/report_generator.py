"""report_generator.py — Generación de reportes .docx en formato SOC oficial."""

import os
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image

# Absolute paths based on script location
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS_DIR = os.path.join(BASE_DIR, "assets")


VERDICT_LABELS = {
    "malicious":  "Alto",
    "suspicious": "Medio",
    "clean":      "Bajo",
    "unknown":    "Desconocido",
}

CRITICIDAD_COLORS = {
    "malicious":  "FF0000",   # Red
    "suspicious": "FF8C00",   # Orange
    "clean":      "008000",   # Green
    "unknown":    "808080",   # Gray
}


def _set_cell_shading(cell, hex_color: str):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders. Values: 'single', 'nil', or None to skip."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="4" w:space="0" w:color="000000"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0):
    """Set internal cell margins (padding) in twips (twentieths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10,
                   color: str = None, alignment: int = None, font_name: str = "Calibri"):
    """Set text in a table cell with formatting. Clears existing content."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_paragraph_to_cell(cell, text: str, bold: bool = False, size: int = 10,
                           color: str = None, alignment: int = None, font_name: str = "Calibri",
                           space_after: int = 2):
    """Add a new paragraph to a cell (preserving existing content)."""
    p = cell.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_bullet_to_cell(cell, text: str, bold_prefix: str = None, size: int = 10,
                        font_name: str = "Calibri"):
    """Add a bullet point paragraph to a cell."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE
    # Add bullet character
    run = p.add_run("•  ")
    run.font.size = Pt(size)
    run.font.name = font_name
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.size = Pt(size)
        run_b.font.name = font_name
        run_b.bold = True
        run_t = p.add_run(text)
        run_t.font.size = Pt(size)
        run_t.font.name = font_name
    else:
        run_t = p.add_run(text)
        run_t.font.size = Pt(size)
        run_t.font.name = font_name
    return p


def _set_table_borders(table):
    """Set all borders for a table to 'single'."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def _merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a given row from col_start to col_end (inclusive)."""
    row = table.rows[row_idx]
    cell = row.cells[col_start]
    for i in range(col_start + 1, col_end + 1):
        cell = cell.merge(row.cells[i])
    return cell


def _get_worst_verdict(results: dict) -> str:
    """Determine the worst verdict across all sources."""
    priority = {"malicious": 4, "suspicious": 3, "unknown": 2, "clean": 1}
    worst = "clean"
    worst_score = 0
    for source_name, result in results.items():
        verdict = result.get("verdict", "unknown").lower()
        score = priority.get(verdict, 0)
        if score > worst_score:
            worst_score = score
            worst = verdict
    return worst


def _build_analysis_text(ioc: str, ioc_type: str, results: dict) -> str:
    """Generate 3-4 sentences of analysis from API results."""
    sentences = []
    malicious_sources = []
    suspicious_sources = []
    clean_sources = []
    countries = set()
    asns = set()
    threat_types = set()
    details = []

    for source_name, result in results.items():
        verdict = result.get("verdict", "unknown").lower()
        detail = result.get("detail", "")
        if verdict == "malicious":
            malicious_sources.append(source_name)
        elif verdict == "suspicious":
            suspicious_sources.append(source_name)
        elif verdict == "clean":
            clean_sources.append(source_name)

        # Extract country info from detail text
        if "país" in detail.lower() or "country" in detail.lower():
            # Try to extract country name
            for word in detail.split():
                if word.istitle() and len(word) > 2 and word.lower() not in ("the", "this", "that"):
                    if word not in ("Malicious", "Suspicious", "Clean", "Unknown"):
                        countries.add(word)

        if detail:
            details.append(f"{source_name}: {detail}")

    # Sentence 1: What was analyzed
    ioc_type_label = {"ip": "IP", "domain": "dominio", "url": "URL", "hash": "hash", "md5": "hash", "sha1": "hash", "sha256": "hash"}.get(ioc_type, ioc_type)
    sentences.append(
        f"Se realizó el análisis del {ioc_type_label} {ioc} utilizando "
        f"múltiples fuentes de inteligencia de amenazas."
    )

    # Sentence 2: Results from sources
    verdict_counts = []
    if malicious_sources:
        verdict_counts.append(f"{len(malicious_sources)} fuente(s) lo clasificaron como malicioso")
    if suspicious_sources:
        verdict_counts.append(f"{len(suspicious_sources)} fuente(s) como sospechoso")
    if clean_sources:
        verdict_counts.append(f"{len(clean_sources)} fuente(s) como limpio")

    if verdict_counts:
        sentences.append(
            "Los resultados de las consultas indican que " + ", ".join(verdict_counts) + "."
        )

    # Sentence 3: Specific threat details
    if malicious_sources:
        sources_str = ", ".join(malicious_sources[:3])
        sentences.append(
            f"Las fuentes {sources_str} reportaron actividad maliciosa asociada al indicador, "
            f"lo que sugiere compromiso o participación en actividades de ciberamenazas."
        )
    elif suspicious_sources:
        sources_str = ", ".join(suspicious_sources[:3])
        sentences.append(
            f"Las fuentes {sources_str} reportaron actividad sospechosa, "
            f"lo que requiere atención del analista para determinar su naturaleza."
        )

    # Additional detail if available
    if details:
        sentences.append(
            "A continuación se presentan los detalles técnicos recopilados durante el análisis."
        )

    return " ".join(sentences)


def _build_risks(verdict: str, results: dict) -> list:
    """Generate 3-4 risks based on verdict and threat type."""
    risks = []

    if verdict == "malicious":
        risks.append(
            ("Compromiso del sistema: ",
             "El indicador ha sido clasificado como malicioso, lo que representa un alto riesgo de compromiso para los activos de la organización.")
        )
        risks.append(
            ("Propagación de malware: ",
             "Posibilidad de que el indicador esté asociado a la distribución de código malicioso, afectando la integridad de los sistemas.")
        )
        risks.append(
            ("Filtración de información: ",
             "Riesgo de exfiltración de datos confidenciales si el indicador está relacionado con canales de comando y control (C2).")
        )
        risks.append(
            ("Acceso no autorizado: ",
             "Potencial acceso remoto no autorizado a la infraestructura, permitiendo movilidad lateral al atacante.")
        )
    elif verdict == "suspicious":
        risks.append(
            ("Actividad anómala: ",
             "El indicador presenta comportamiento sospechoso que podría ser indicativo de una amenaza en etapa temprana.")
        )
        risks.append(
            ("Falso positivo potencial: ",
             "Se requiere verificación adicional para descartar que se trate de un servicio legítimo mal clasificado.")
        )
        risks.append(
            ("Escalada de privilegios: ",
             "Si se confirma la actividad maliciosa, el atacante podría intentar escalar privilegios en los sistemas afectados.")
        )
        risks.append(
            ("Persistencia: ",
             "Posibilidad de que el indicador esté asociado a mecanismos de persistencia en la red.")
        )
    else:  # clean or unknown
        risks.append(
            ("Bajo riesgo: ",
             "El indicador no presenta actividad maliciosa confirmada según las fuentes consultadas.")
        )
        risks.append(
            ("Falso negativo potencial: ",
             "Aunque no se detectaron amenazas, se recomienda monitoreo continuo para detectar cambios en la reputación del indicador.")
        )
        risks.append(
            ("Contexto faltante: ",
             "La falta de detecciones no garantiza que el indicador sea benigno; debe evaluarse en el contexto de la infraestructura del cliente.")
        )
        risks.append(
            ("Recomendación de monitoreo: ",
             "Se sugiere mantener el indicador en observación y repetir el análisis periódicamente.")
        )

    return risks


def _build_ioc_data_table_text(ioc: str, ioc_type: str, results: dict) -> list:
    """Build IOC data for the sub-table rows in ANÁLISIS section."""
    rows = []
    # Gather some detail info
    country = ""
    asn = ""
    threat_type = ""
    for source_name, result in results.items():
        detail = result.get("detail", "")
        result_verdict = result.get("verdict", "")
        if not country:
            # Try to extract basic info from details
            pass

    # Pre-fill with available data, leaving rest blank for analyst
    rows.append({
        "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "accion": "En análisis",
        "ip_origen": ioc if ioc_type == "ip" else "",
        "ip_destino": "",
        "puerto_dest": "",
        "rule": "",
        "zona": ""
    })
    return rows


def generate_word_report(ioc: str, ioc_type: str, results: dict, ai_summary: str,
                         screenshots_dict: dict) -> bytes:
    """
    Generate a .docx report for a single IoC in the official SOC format.

    Args:
        ioc: The indicator of compromise string.
        ioc_type: Type of IoC (ip, domain, url, hash).
        results: Dictionary mapping source_name -> {verdict, detail, web_url, ...}.
        ai_summary: Executive summary text from AI (Groq).
        screenshots_dict: Dictionary mapping source_name -> PNG bytes (or None).

    Returns:
        .docx file as bytes.
    """
    doc = Document()

    # -- Default style --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

    # Narrow margins for more space
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # -- Determine worst verdict --
    worst_verdict = _get_worst_verdict(results)
    criticidad_label = VERDICT_LABELS.get(worst_verdict, "Desconocido")
    criticidad_color = CRITICIDAD_COLORS.get(worst_verdict, "808080")

    # ===================================================================
    # PAGE 1 — HEADER: Three-column table
    # ===================================================================
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(header_table)

    # Left cell: Axity logo
    cell_left = header_table.rows[0].cells[0]
    cell_left.text = ""
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    axity_path = os.path.join(ASSETS_DIR, "Axity_Logo.png")
    run_left = p_left.add_run()
    run_left.add_picture(axity_path, width=Cm(4))

    # Center cell: Title
    cell_center = header_table.rows[0].cells[1]
    _set_cell_text(cell_center, "Reporte de Evento", bold=True, size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph_to_cell(cell_center, "CORTEX XDR", bold=False, size=11,
                           color="555555", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    # Right cell: Cortex logo (converted from WebP to PNG in memory)
    cell_right = header_table.rows[0].cells[2]
    cell_right.text = ""
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cortex_webp_path = os.path.join(ASSETS_DIR, "Cortex-logo.webp")
    img = Image.open(cortex_webp_path)
    cortex_buf = BytesIO()
    img.save(cortex_buf, format="PNG")
    cortex_buf.seek(0)
    run_right = p_right.add_run()
    run_right.add_picture(cortex_buf, width=Cm(4))

    # Set header column widths
    for row in header_table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(8)
        row.cells[2].width = Cm(4.5)

    doc.add_paragraph()  # spacer

    # ===================================================================
    # PAGE 1 — METADATA TABLE
    # ===================================================================
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(meta_table)

    # Row 1: EVENTO
    _set_cell_text(meta_table.rows[0].cells[0], "EVENTO", bold=True, size=10)
    _set_cell_text(meta_table.rows[0].cells[1], "", bold=False, size=10)

    # Row 2: FUENTE
    _set_cell_text(meta_table.rows[1].cells[0], "FUENTE", bold=True, size=10)
    _set_cell_text(meta_table.rows[1].cells[1], "", bold=False, size=10)

    # Row 3: CRITICIDAD (colored)
    _set_cell_text(meta_table.rows[2].cells[0], "CRITICIDAD", bold=True, size=10)
    _set_cell_text(meta_table.rows[2].cells[1], criticidad_label, bold=True, size=10,
                   color=criticidad_color)

    # Set metadata column widths
    meta_table.columns[0].width = Cm(3.5)
    meta_table.columns[1].width = Cm(13.5)
    for row in meta_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13.5)

    doc.add_paragraph()  # spacer

    # ===================================================================
    # PAGE 1 — ANÁLISIS SECTION
    # ===================================================================
    analisis_table = doc.add_table(rows=1, cols=2)
    analisis_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(analisis_table)

    # Left cell: label
    cell_label = analisis_table.rows[0].cells[0]
    _set_cell_text(cell_label, "ANÁLISIS", bold=True, size=10,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_label.width = Cm(3.5)
    analisis_table.columns[0].width = Cm(3.5)

    # Right cell: content
    cell_content = analisis_table.rows[0].cells[1]
    cell_content.width = Cm(13.5)
    analisis_table.columns[1].width = Cm(13.5)
    # Clear default paragraph
    cell_content.text = ""

    # Analysis text (3-4 sentences)
    analysis_text = _build_analysis_text(ioc, ioc_type, results)
    _add_paragraph_to_cell(cell_content, analysis_text, size=10, space_after=6)

    # Sub-table inside: Fecha y Hora | Acción | IP Origen | IP Destino
    sub_table = doc.add_table(rows=2, cols=4)
    sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(sub_table)

    # Header row
    sub_headers = ["Fecha y Hora", "Acción", "IP Origen", "IP Destino"]
    for idx, hdr in enumerate(sub_headers):
        _set_cell_text(sub_table.rows[0].cells[idx], hdr, bold=True, size=8,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(sub_table.rows[0].cells[idx], "E0E0E0")
        _set_cell_margins(sub_table.rows[0].cells[idx], top=40, bottom=40, left=60, right=60)

    # Set sub-table column widths: total = 3.5+2.5+3.5+4.0 = 13.5cm
    sub_col_widths = [Cm(3.5), Cm(2.5), Cm(3.5), Cm(4.0)]
    for row in sub_table.rows:
        for idx in range(4):
            row.cells[idx].width = sub_col_widths[idx]
    for idx in range(4):
        sub_table.columns[idx].width = sub_col_widths[idx]

    # Data row (pre-fill what's available)
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    data_values = [
        now_str,
        "En análisis",
        ioc if ioc_type == "ip" else "",
        ""
    ]
    for idx, val in enumerate(data_values):
        _set_cell_text(sub_table.rows[1].cells[idx], val, size=8,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_margins(sub_table.rows[1].cells[idx], top=40, bottom=40, left=60, right=60)

    # Add sub-table as inline content in the cell
    # We need to move the sub_table into cell_content
    # python-docx doesn't natively support nested tables, but we can use XML manipulation
    # Remove sub_table from document body and add to cell
    cell_content._tc.append(sub_table._tbl)

    # 2 closing sentences
    closing = (
        "Se recomienda correlacionar esta información con otros eventos de seguridad "
        "en la infraestructura del cliente para determinar el alcance real de la amenaza. "
        "El presente análisis debe ser complementado con la experiencia del analista asignado."
    )
    _add_paragraph_to_cell(cell_content, closing, size=10, space_after=4)

    # Space for analyst
    _add_paragraph_to_cell(cell_content, "", size=10, space_after=2)
    _add_paragraph_to_cell(cell_content,
        "--- Espacio para notas del analista ---", size=9,
        color="999999", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    doc.add_paragraph()  # spacer

    # ===================================================================
    # PAGE 1 — RIESGOS IDENTIFICADOS SECTION
    # ===================================================================
    riesgos_table = doc.add_table(rows=1, cols=2)
    riesgos_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(riesgos_table)

    # Left cell: label
    cell_riesgos_label = riesgos_table.rows[0].cells[0]
    _set_cell_text(cell_riesgos_label, "RIESGOS IDENTIFICADOS", bold=True, size=10,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_riesgos_label.width = Cm(3.5)
    riesgos_table.columns[0].width = Cm(3.5)

    # Right cell: bullet list
    cell_riesgos = riesgos_table.rows[0].cells[1]
    cell_riesgos.width = Cm(13.5)
    riesgos_table.columns[1].width = Cm(13.5)
    cell_riesgos.text = ""

    risks = _build_risks(worst_verdict, results)
    for title, desc in risks:
        _add_bullet_to_cell(cell_riesgos, desc, bold_prefix=title, size=10)

    # ===================================================================
    # PAGE BREAK → PAGE 2
    # ===================================================================
    doc.add_page_break()

    # ===================================================================
    # PAGE 2 — IOC DATA TABLE
    # ===================================================================
    ioc_headers = ["DIR IP ORIGEN", "PUERTO ORIGEN", "DIR IP DESTINO", "PUERTO DESTINO"]
    ioc_table = doc.add_table(rows=2, cols=4)
    ioc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(ioc_table)

    for idx, hdr in enumerate(ioc_headers):
        _set_cell_text(ioc_table.rows[0].cells[idx], hdr, bold=True, size=10,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(ioc_table.rows[0].cells[idx], "E0E0E0")

    # Pre-fill DIR IP ORIGEN with the analyzed IoC
    _set_cell_text(ioc_table.rows[1].cells[0], ioc if ioc_type == "ip" else ioc,
                   size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for idx in range(1, 4):
        _set_cell_text(ioc_table.rows[1].cells[idx], "", size=10,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()  # spacer

    # ===================================================================
    # PAGE 2 — RECOMENDACIONES SECTION
    # ===================================================================
    reco_table = doc.add_table(rows=1, cols=2)
    reco_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(reco_table)

    cell_reco_label = reco_table.rows[0].cells[0]
    _set_cell_text(cell_reco_label, "RECOMENDACIONES", bold=True, size=10,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_reco_label.width = Cm(3.5)
    reco_table.columns[0].width = Cm(3.5)

    cell_reco = reco_table.rows[0].cells[1]
    cell_reco.width = Cm(13.5)
    reco_table.columns[1].width = Cm(13.5)
    cell_reco.text = ""

    if ai_summary:
        # Split AI summary into bullet points by sentences
        import re
        sentences = re.split(r'(?<=[.!])\s+', ai_summary.strip())
        for sent in sentences:
            if sent.strip():
                _add_bullet_to_cell(cell_reco, sent.strip(), size=10)
    else:
        _add_paragraph_to_cell(cell_reco, "No disponible.", size=10)

    doc.add_paragraph()  # spacer

    # ===================================================================
    # PAGE 2 — EVIDENCIA SECTION
    # ===================================================================
    evidencia_table = doc.add_table(rows=1, cols=2)
    evidencia_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(evidencia_table)

    # Left cell: label spanning vertically
    cell_evidencia_label = evidencia_table.rows[0].cells[0]
    _set_cell_text(cell_evidencia_label, "EVIDENCIA", bold=True, size=10,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_evidencia_label.width = Cm(3.5)
    evidencia_table.columns[0].width = Cm(3.5)

    # Right cell: screenshots
    cell_evidencia = evidencia_table.rows[0].cells[1]
    cell_evidencia.width = Cm(13.5)
    evidencia_table.columns[1].width = Cm(13.5)
    cell_evidencia.text = ""

    # Title with IoC type
    ioc_type_spanish = {"ip": "IP", "domain": "Dominio", "url": "URL", "hash": "Hash", "md5": "Hash", "sha1": "Hash", "sha256": "Hash"}.get(ioc_type, ioc_type)
    _add_paragraph_to_cell(cell_evidencia, f"Categorización {ioc_type_spanish}:",
                           bold=True, size=11, space_after=6)

    # Add screenshots
    screenshots_added = False
    for source_name in sorted(screenshots_dict.keys()):
        screenshot_bytes = screenshots_dict[source_name]
        if screenshot_bytes:
            screenshots_added = True
            # Source name as caption
            _add_paragraph_to_cell(cell_evidencia, source_name, bold=True, size=10,
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
            # Embed image
            image_stream = BytesIO(screenshot_bytes)
            p_img = cell_evidencia.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(image_stream, width=Cm(12.5))
            # Spacer after image
            sp = cell_evidencia.add_paragraph()
            sp.paragraph_format.space_after = Pt(4)
            sp.paragraph_format.space_before = Pt(0)

    if not screenshots_added:
        _add_paragraph_to_cell(cell_evidencia, "No se generaron capturas de evidencia.",
                               size=10, color="808080")

    doc.add_paragraph()  # spacer

    # ===================================================================
    # PAGE 2 — FOOTER TABLE
    # ===================================================================
    footer_table = doc.add_table(rows=2, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(footer_table)

    # ANALISTA row
    _set_cell_text(footer_table.rows[0].cells[0], "ANALISTA", bold=True, size=10)
    _set_cell_text(footer_table.rows[0].cells[1], "", size=10)  # blank for manual fill
    footer_table.rows[0].cells[0].width = Cm(3.5)
    footer_table.rows[0].cells[1].width = Cm(13.5)

    # FECHA DEL ANÁLISIS row
    now_date = datetime.now().strftime("%Y-%m-%d %H:%M")
    _set_cell_text(footer_table.rows[1].cells[0], "FECHA DEL ANÁLISIS", bold=True, size=10)
    _set_cell_text(footer_table.rows[1].cells[1], now_date, size=10)
    footer_table.rows[1].cells[0].width = Cm(3.5)
    footer_table.rows[1].cells[1].width = Cm(13.5)

    # Set footer column widths
    footer_table.columns[0].width = Cm(3.5)
    footer_table.columns[1].width = Cm(13.5)

    doc.add_paragraph()  # spacer

    # ===================================================================
    # FINAL NOTE (outside table, italic small text)
    # ===================================================================
    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_para.paragraph_format.space_before = Pt(6)
    note_run = note_para.add_run(
        "Nota: El nivel de alerta presentado en este informe es una evaluación subjetiva "
        "realizada tanto por la herramienta (Cortex XDR) como por los analistas de axity. "
        "La importancia de cada alerta de vulnerabilidad puede variar dependiendo del contexto "
        "específico de la infraestructura del cliente, así como de los activos y accesos "
        "involucrados. Agradecemos cualquier retroalimentación adicional sobre este informe, "
        "ya que nos ayuda a mejorar continuamente nuestros procesos de análisis."
    )
    note_run.font.size = Pt(8)
    note_run.font.name = "Calibri"
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # -- Footer with page number --
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("IOC Enricher — Uso interno SOC  |  Pág. ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = "Calibri"
        # Page number field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_fld = p.add_run()
        run_fld._r.append(fld_char_begin)
        run_page = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_page._r.append(instr)
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_end = p.add_run()
        run_end._r.append(fld_char_end)

    # -- Save to bytes --
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()